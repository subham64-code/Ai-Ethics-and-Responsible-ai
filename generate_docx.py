import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ===== STYLES =====
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.15

# ===== TITLE =====
title = doc.add_heading('Chapter 4: Ethics and Responsible AI', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Class 8 - Computational Thinking and Artificial Intelligence')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)
run.italic = True

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub2.add_run('CBSE Student Handbook | Simplified & Enriched Edition')
run2.font.size = Pt(12)
run2.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()  # spacer

# ===== IMAGE 1: Chapter Header =====
img_path = r'D:\open\extracted_images\ch4_page69_img1.jpeg'
if os.path.exists(img_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(5.5))
    doc.add_paragraph()

# ===== INTRODUCTION =====
doc.add_heading('What is this Chapter About?', level=1)
doc.add_paragraph(
    'In earlier chapters, you learned how Artificial Intelligence (AI) systems collect data, '
    'recognize patterns, and make predictions. AI helps us in education, healthcare, business, '
    'and daily life. But as AI becomes more powerful, we must ask:'
)
p = doc.add_paragraph()
run = p.add_run('"How should AI be created and used responsibly?"')
run.bold = True
run.font.size = Pt(13)

p2 = doc.add_paragraph()
run2 = p2.add_run('Responsible AI')
run2.bold = True
p2.add_run(' means designing and using artificial intelligence in ways that are ')
run3 = p2.add_run('fair, safe, transparent, and beneficial to society.')
run3.bold = True

doc.add_paragraph(
    'Technology does not exist alone. It affects people, communities, and society. '
    'The choices AI systems make can influence opportunities, decisions, and even people\'s lives. '
    'That is why technology must be guided by ethical principles.'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Key Questions We Will Explore:')
run.bold = True

questions = [
    '1. What do we mean by AI ethics?',
    '2. Why is privacy important in the digital world?',
    '3. How can bias affect AI decisions?',
    '4. What is misinformation, and why does it spread so easily?',
    '5. Who is responsible when AI systems make mistakes?'
]
for q in questions:
    p = doc.add_paragraph(q, style='List Bullet')

# ===== SECTION 4.1 =====
doc.add_page_break()
doc.add_heading('4.1 What is AI Ethics?', level=1)

p = doc.add_paragraph()
run = p.add_run('When new technology is created, people often ask: "What can this technology do?" ')
p.add_run('But there is another question that is equally important: ')
run2 = p.add_run('"What should this technology do?"')
run2.bold = True
p.add_run(' This second question takes us into the world of ethics.')

doc.add_heading('Definition', level=2)
p = doc.add_paragraph()
run = p.add_run('AI Ethics')
run.bold = True
p.add_run(' refers to the values and principles that guide how artificial intelligence systems should be designed, developed, and used.')

doc.add_heading('Why Ethics Matters', level=2)
p = doc.add_paragraph('Ethics remind us that technology must serve people and society. AI ethics ensure that systems:')

items = [
    ('Fair', 'Treat everyone equally'),
    ('Respect privacy', 'Protect personal information'),
    ('Help society', 'Benefit people, not harm them'),
    ('Remain accountable', 'Someone is responsible for their actions')
]

for title_text, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{title_text} — {desc}')

doc.add_heading('Simple Example', level=2)
doc.add_paragraph(
    'Imagine an AI system used in a hospital to help doctors. Ethical AI would:\n'
    '• Treat all patients equally (rich or poor)\n'
    '• Keep patient data private\n'
    '• Help doctors make better decisions\n'
    '• Allow doctors to overrule the AI if needed'
)

p = doc.add_paragraph()
run = p.add_run('Reflect: ')
run.bold = True
run.font.color.rgb = RGBColor(0, 100, 0)
p.add_run('Can a system be considered successful if it is fast and accurate but harms certain groups of people?')

# ===== SECTION 4.2 =====
doc.add_page_break()
doc.add_heading('4.2 Privacy in the Digital World', level=1)

doc.add_heading('What is Privacy?', level=2)
p = doc.add_paragraph()
run = p.add_run('Privacy')
run.bold = True
p.add_run(' means keeping your personal information safe and under your control.')

doc.add_heading('What Data Do Apps Collect?', level=2)
table = doc.add_table(rows=7, cols=2, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Type of Data', 'Example']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

data = [
    ('Your name', 'Full name, username'),
    ('Your location', 'GPS coordinates, city'),
    ('Your photographs', 'Selfies, camera roll'),
    ('Health information', 'Step count, heart rate'),
    ('Browsing history', 'Websites you visit'),
    ('Contact details', 'Phone numbers, email')
]
for row_idx, (col1, col2) in enumerate(data, 1):
    table.rows[row_idx].cells[0].text = col1
    table.rows[row_idx].cells[1].text = col2

doc.add_paragraph()

doc.add_heading('Why is Privacy Important?', level=2)
doc.add_paragraph('If personal information is misused, it can cause serious problems:\n'
    '• Identity theft — Someone pretends to be you\n'
    '• Financial fraud — Money stolen from your accounts\n'
    '• Reputation damage — Embarrassing information shared publicly')

doc.add_heading('Before Sharing Data Online, Ask:', level=2)
items2 = ['Why is this information needed?', 'Who will be able to see it?',
          'How long will it be stored?', 'Can I delete it later?']
for item in items2:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('Good Digital Habits: ')
run.bold = True
p.add_run('Never share sensitive information like OTPs, bank account details, passwords, or PINs.')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Remember: ')
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 180)
p.add_run('Privacy is NOT about hiding from the world. It is about having control over your personal information.')

# ===== IMAGE 2 =====
img2 = r'D:\open\extracted_images\ch4_page71_img1.jpeg'
if os.path.exists(img2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = p.add_run('\n[App Permission Scenario: Always think before clicking Allow]')
    caption.font.size = Pt(9)
    caption.italic = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run()
    run.add_picture(img2, width=Inches(4.5))

# ===== SECTION 4.3 =====
doc.add_page_break()
doc.add_heading('4.3 Bias and Fairness in AI', level=1)

doc.add_heading('What is Bias in AI?', level=2)
doc.add_paragraph(
    'Artificial Intelligence systems learn from data. But what happens if the data itself is '
    'incomplete or unfair? The system may learn patterns that are also unfair. This is known as bias.'
)

doc.add_heading('Problems Caused by AI Bias', level=2)
problems = [
    'Reject deserving candidates — qualified people not getting jobs',
    'Spread stereotypes — associating certain jobs with specific genders',
    'Ignore minority groups — not recognizing faces of certain ethnicities',
    'Make incorrect predictions — wrong medical diagnoses'
]
for prob in problems:
    doc.add_paragraph(prob, style='List Bullet')

doc.add_heading('Real-World Example', level=2)
doc.add_paragraph(
    'A facial recognition system trained mostly with images of people from one region may not '
    'correctly recognize people from other regions. In a famous case, an AI system could not '
    'detect a dark-skinned person\'s face until she wore a white mask!'
)

doc.add_heading('How Can Bias Be Reduced?', level=2)
table2 = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Strategy', 'How It Helps']):
    table2.rows[0].cells[i].text = h
    for paragraph in table2.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

strategies = [
    ('Use diverse data', 'Include data from all groups'),
    ('Test carefully', 'Check before releasing'),
    ('Human review', 'Let people check AI decisions'),
    ('Monitor regularly', 'Keep checking performance')
]
for idx, (s1, s2) in enumerate(strategies, 1):
    table2.rows[idx].cells[0].text = s1
    table2.rows[idx].cells[1].text = s2

# ===== SECTION 4.4 =====
doc.add_page_break()
doc.add_heading('4.4 Misinformation and Social Impact', level=1)

doc.add_heading('What is Misinformation?', level=2)
p = doc.add_paragraph()
run = p.add_run('Misinformation')
run.bold = True
p.add_run(' means incorrect or misleading information shared with others.')

doc.add_heading('How AI Makes It Worse', level=2)
doc.add_paragraph(
    'Modern AI tools can:\n'
    '• Create realistic fake images\n'
    '• Generate fake news articles\n'
    '• Imitate voices (voice cloning)\n'
    '• Produce misleading videos (Deepfakes)'
)

doc.add_heading('Impact of Misinformation', level=2)
doc.add_paragraph(
    'Misinformation can affect:\n'
    '• Public opinion — People believe false things\n'
    '• Elections — Fake news influences voters\n'
    '• Careers and reputations — False accusations\n'
    '• Public trust — People stop believing anything\n'
    '• Social harmony — Creates divisions'
)

doc.add_heading('Before Believing or Forwarding:', level=2)
tips = ['Check the source — Who created it?',
        'Verify facts — Use trusted websites',
        'Read beyond headlines',
        'Think critically — Does it make sense?']
for t in tips:
    doc.add_paragraph(t, style='List Bullet')

# ===== SECTION 4.5 =====
doc.add_page_break()
doc.add_heading('4.5 Accountability and Human Control', level=1)

doc.add_heading('What is Accountability?', level=2)
p = doc.add_paragraph()
run = p.add_run('Accountability')
run.bold = True
p.add_run(' means that humans remain responsible for the decisions and actions of AI systems.')

doc.add_heading('Important Questions', level=2)
for q in ['Who created the system?', 'Who checks whether it works correctly?',
          'Who fixes mistakes if they occur?', 'Who is responsible if harm happens?']:
    doc.add_paragraph(q, style='List Bullet')

doc.add_paragraph(
    'AI should assist human decision-making, not replace human judgment. '
    'Machines can analyze data quickly, but final responsibility must always remain with people.'
)

p = doc.add_paragraph()
run = p.add_run('Key Fields Where Human Supervision is Essential: ')
run.bold = True
doc.add_paragraph('Healthcare, Education, Banking, Law, Public services', style='List Bullet')

# ===== IMAGE 3 =====
img3 = r'D:\open\extracted_images\ch4_page73_img1.jpeg'
if os.path.exists(img3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = p.add_run('\n[Human Accountability: Doctors reviewing AI recommendations]')
    caption.font.size = Pt(9)
    caption.italic = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run()
    run.add_picture(img3, width=Inches(3.5))

# ===== SUMMARY TABLE =====
doc.add_page_break()
doc.add_heading('Chapter Summary', level=1)
table3 = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
summary_data = [
    ('Concept', 'Key Point'),
    ('AI Ethics', 'Values and principles guiding responsible AI design'),
    ('Privacy', 'Protecting and controlling personal data'),
    ('Bias & Fairness', 'Unfair data leads to unfair AI outcomes'),
    ('Misinformation', 'False information spreads easily with digital tools'),
    ('Accountability', 'Humans must remain responsible for AI decisions')
]
for idx, (c1, c2) in enumerate(summary_data):
    table3.rows[idx].cells[0].text = c1
    table3.rows[idx].cells[1].text = c2
    if idx == 0:
        for cell in table3.rows[idx].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

# ===== MATH INTEGRATION =====
doc.add_page_break()
doc.add_heading('Mathematics Integration: Comparing Quantities with AI', level=1)

doc.add_paragraph(
    'Just as AI compares data to make decisions, we compare quantities in mathematics. '
    'Let us see how ratios and percentages (from NCERT Class 8, Chapter 8) relate to understanding AI fairness.'
)

doc.add_heading('Ratio Concepts', level=2)
doc.add_paragraph(
    'A ratio compares two quantities. For example, in the sports recommendation example:\n'
    '• Ratio of boys who like Cricket to Girls who like Cricket = 55 : 15 = 11 : 3\n'
    '• This shows a biased dataset — the AI learned an unfair pattern!'
)

doc.add_heading('Percentage Concepts', level=2)
doc.add_paragraph(
    'Percentage tells us how many out of 100.\n\n'
    'Example: If an AI training dataset has 80 images of cats and 20 images of dogs:\n'
    '• Percentage of cat images = (80/100) x 100 = 80%\n'
    '• Percentage of dog images = (20/100) x 100 = 20%\n'
    '• The AI will be better at recognizing cats because of more data!'
)

doc.add_heading('Practice Problems', level=2)

problems_data = [
    '1. An AI training dataset has 120 images of apples and 80 images of oranges.\n   (a) Find the ratio of apple images to orange images.\n   (b) What percentage of images are apples?\n   (c) What percentage of images are oranges?\n\n'
    '2. In a school, an AI system recommends books. It was trained on data where 150 boys and 50 girls read science fiction.\n   (a) Find the ratio of boys to girls who read science fiction.\n   (b) What percentage of science fiction readers are boys?\n\n'
    '3. A facial recognition system was trained with 1000 photos. 800 photos are of light-skinned people and 200 are of dark-skinned people.\n   (a) Find the percentage of light-skinned training photos.\n   (b) Find the percentage of dark-skinned training photos.\n   (c) Why might this AI system be biased?\n\n'
    '4. A company spent Rs 50,000 developing an AI system that saves Rs 8,000 per month.\n   (a) How many months to recover the investment?\n   (b) What is the savings as a percentage of the investment after 12 months?\n\n'
    '5. An AI software costs Rs 15,000 but is offered at a 20% discount.\n   (a) Find the discount amount.\n   (b) Find the selling price.'
]

for prob in problems_data:
    p = doc.add_paragraph(prob)

# ===== ACTIVITIES =====
doc.add_page_break()
doc.add_heading('Activities', level=1)

doc.add_heading('Activity 1: Privacy Check', level=2)
doc.add_paragraph(
    'Objective: Understand what data apps collect\n\n'
    'Steps:\n'
    '1. Look at 3 apps on your phone or computer\n'
    '2. Check what permissions they ask for (camera, location, contacts, etc.)\n'
    '3. For each app, ask: "Why does this app need this permission?"\n'
    '4. Create a table with: App Name | Permissions Asked | Is it Needed?'
)

doc.add_heading('Activity 2: Bias Detection', level=2)
doc.add_paragraph(
    'Objective: Identify bias in data\n\n'
    'Scenario: You are building an AI system to recommend movies. Your training data has:\n'
    'Action: 60 boys, 10 girls | Comedy: 20 boys, 40 girls | Romance: 5 boys, 35 girls\n\n'
    'Questions:\n'
    '1. What movie type will the AI recommend to a boy? A girl?\n'
    '2. Is this recommendation system fair?\n'
    '3. How could you make the dataset more balanced?'
)

doc.add_heading('Activity 3: Create a Fair Dataset', level=2)
doc.add_paragraph(
    'Collect 20 images of objects from home (5 of each type - spoons, pens, books, cups).\n'
    'Ask a friend to test if your dataset is balanced. Would an AI trained on this data be fair?'
)

doc.add_heading('Activity 4: Misinformation Detective', level=2)
doc.add_paragraph(
    'Find 2 news headlines and answer:\n'
    '• Who wrote it? Is the source trustworthy?\n'
    '• Can you find the same news on a different website?\n'
    '• Are there emotional words trying to influence you?'
)

doc.add_heading('Activity 5: AI Ethics Poster', level=2)
doc.add_paragraph(
    'Create a poster showing:\n'
    '• One key message about AI ethics\n'
    '• A simple example\n'
    '• A colorful illustration\n'
    'Display it in your classroom or share online.'
)

# ===== EXERCISE =====
doc.add_page_break()
doc.add_heading('Exercise Questions', level=1)

doc.add_heading('A. Multiple Choice Questions', level=2)
mcqs = [
    '1. AI ethics focus on:\n   a) Making machines faster\n   b) Making machines cheaper\n   c) Ensuring responsible use\n   d) Deleting apps',
    '2. Privacy means:\n   a) Sharing everything online\n   b) Protecting personal information\n   c) Hiding from society\n   d) Deleting apps',
    '3. AI bias can occur when:\n   a) AI systems are turned off\n   b) Data used for training is incomplete or unfair\n   c) Computers are slow\n   d) The internet is not working',
    '4. Before sharing information online, what should you do?\n   a) Share it quickly\n   b) Think about why the information is needed\n   c) Forward it to everyone\n   d) Ignore the message',
    '5. AI systems learn patterns from:\n   a) Books\n   b) Data\n   c) Games\n   d) Images'
]
for mcq in mcqs:
    doc.add_paragraph(mcq)

doc.add_heading('B. Fill in the Blanks', level=2)
fill_blanks = [
    '1. Misinformation means _____________ information.',
    '2. Fair AI systems treat people _____________ .',
    '3. Humans must remain _____________ for AI decisions.',
    '4. AI systems learn patterns from _____________ .',
    '5. Incorrect or misleading information shared online is called _____________ .'
]
for fb in fill_blanks:
    doc.add_paragraph(fb)

doc.add_heading('C. Short Answer Questions', level=2)
short_qs = [
    '1. Why is privacy important in the digital world?',
    '2. Give one example of misinformation.',
    '3. What does accountability mean in AI systems?',
    '4. What are AI ethics?',
    '5. Why is human supervision important when AI systems are used?'
]
for sq in short_qs:
    doc.add_paragraph(sq)

doc.add_heading('D. Case-Based Questions', level=2)
cases = [
    '1. A student installs an app that asks for permission to access location, contacts, and photos. The student pauses and thinks about why the app needs this information.\n   • What important digital habit is the student practising?',
    '2. An AI system used in healthcare suggests that a patient has a high-risk level. Doctors carefully review the AI recommendation before making a final decision.\n   • Why is human review important in this situation?',
    '3. An AI system is trained mostly with data from one city and later used in other places. The system does not work well for people from different regions.\n   • What problem in AI does this example show?',
    '4. A message spreads quickly on social media, but the information in it is incorrect.\n   • What is this situation called, and what should users do before sharing such messages?'
]
for case in cases:
    doc.add_paragraph(case)

# ===== GLOSSARY =====
doc.add_page_break()
doc.add_heading('Glossary', level=1)
glossary = [
    ('AI Ethics', 'Values and principles that guide how AI should be designed and used'),
    ('Privacy', 'Keeping personal information safe and under your control'),
    ('Bias', 'Unfair patterns learned by AI from unbalanced data'),
    ('Fairness', 'Treating all people equally and without prejudice'),
    ('Misinformation', 'False or misleading information shared with others'),
    ('Accountability', 'Being responsible for decisions and actions'),
    ('Responsible AI', 'AI that is fair, safe, transparent, and beneficial'),
    ('Training Data', 'The data used to teach an AI system'),
    ('Deepfake', 'AI-generated fake video or audio that looks real'),
    ('Transparency', 'Being open about how AI systems work')
]

table4 = doc.add_table(rows=len(glossary)+1, cols=2, style='Light Grid Accent 1')
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
table4.rows[0].cells[0].text = 'Term'
table4.rows[0].cells[1].text = 'Meaning'
for paragraph in table4.rows[0].cells[0].paragraphs:
    for run in paragraph.runs:
        run.bold = True
for paragraph in table4.rows[0].cells[1].paragraphs:
    for run in paragraph.runs:
        run.bold = True

for idx, (term, meaning) in enumerate(glossary, 1):
    table4.rows[idx].cells[0].text = term
    table4.rows[idx].cells[1].text = meaning

# ===== PROJECT IDEAS =====
doc.add_page_break()
doc.add_heading('Project Ideas', level=1)
doc.add_heading('Project 1: Design Your Own Fair AI', level=2)
doc.add_paragraph(
    'Create a simple AI model using Teachable Machine (teachablemachine.withgoogle.com) that:\n'
    '• Recognizes at least 3 different objects\n'
    '• Uses balanced training data\n'
    '• Test if it works fairly for different inputs'
)

doc.add_heading('Project 2: AI Ethics Survey', level=2)
doc.add_paragraph(
    'Conduct a survey in your school:\n'
    '• Ask 20 students: "What do you know about AI ethics?"\n'
    '• Ask 20 students: "Have you shared information online without checking?"\n'
    '• Create a report with graphs showing your findings'
)

doc.add_heading('Project 3: Create Awareness Material', level=2)
doc.add_paragraph(
    'Create a short video, presentation, or pamphlet about:\n'
    '• One aspect of AI ethics (privacy, bias, misinformation, or accountability)\n'
    '• Include examples from daily life\n'
    '• Share it with your family or community'
)

# Save
doc.save(r'D:\open\Chapter-4_Ethics_and_Responsible_AI.docx')
print("DOCX file created successfully!")

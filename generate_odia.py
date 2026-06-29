import os, sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
style = doc.styles["Normal"]
font = style.font
font.name = "Nirmala"
font.size = Pt(12)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Nirmala")
style.paragraph_format.line_spacing = 1.3

def add_head(text, lv=1):
    h = doc.add_heading(text, level=lv)
    for r in h.runs:
        r.font.name = "Nirmala"
        r.element.rPr.rFonts.set(qn("w:eastAsia"), "Nirmala")
    return h

def add_p(text, bold=False, size=12, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Nirmala"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Nirmala")
    run.font.size = Pt(size)
    if bold: run.bold = True
    if color: run.font.color.rgb = color
    return p

def add_b(text, bp=None):
    p = doc.add_paragraph(style="List Bullet")
    if bp:
        r = p.add_run(bp)
        r.bold = True
        r.font.name = "Nirmala"
        r.element.rPr.rFonts.set(qn("w:eastAsia"), "Nirmala")
    r2 = p.add_run(text)
    r2.font.name = "Nirmala"
    r2.element.rPr.rFonts.set(qn("w:eastAsia"), "Nirmala")
    return p

def add_img(path, w=5.0, cap=None):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(w))
        if cap:
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p2.add_run(cap)
            r.font.name = "Nirmala"
            r.element.rPr.rFonts.set(qn("w:eastAsia"), "Nirmala")
            r.font.size = Pt(9)
            r.italic = True

def add_tab(headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers), style="Light Grid Accent 1")
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True; r.font.name = "Nirmala"
                r.element.rPr.rFonts.set(qn("w:eastAsia"), "Nirmala")
    for ri, rd in enumerate(rows, 1):
        for ci, v in enumerate(rd):
            c = t.rows[ri].cells[ci]; c.text = v
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = "Nirmala"
                    r.element.rPr.rFonts.set(qn("w:eastAsia"), "Nirmala")

print("Functions defined OK")

# ===== TITLE PAGE =====
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Chapter 4 \u0B06\u0B30\u0B4D\u0B1F\u0B3F\u0B15\u0B3F\u0B38\u0B3F\u0B06\u0B32\u0B4D \u0B07\u0B23\u0B4D\u0B1F\u0B47\u0B32\u0B3F\u0B1C\u0B47\u0B28\u0B4D\u0B38 \u2013 \u0B28\u0B40\u0B24\u0B3F\u0B36\u0B3E\u0B38\u0B4D\u0B24\u0B4D\u0B30 \u0B0F\u0B2C\u0B02 \u0B26\u0B3E\u0B2F\u0B3F\u0B24\u0B4D\u0B71\u0B2A\u0B42\u0B30\u0B4D\u0B23 AI")
r.font.name = "Nirmala"
r.element.rPr.rFonts.set(qn("w:eastAsia"), "Nirmala")
r.font.size = Pt(26)
r.bold = True
r.font.color.rgb = RGBColor(0, 51, 102)

# Simplified title for class 3
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Class 3 \u0B2A\u0B3E\u0B07\u0B02 \u0B38\u0B30\u0B33 \u0B2C\u0B4D\u0B2F\u0B3E\u0B16\u0B4D\u0B2F\u0B3E")
r2.font.name = "Nirmala"
r2.element.rPr.rFonts.set(qn("w:eastAsia"), "Nirmala")
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(100, 100, 100)

print("Title page OK")
